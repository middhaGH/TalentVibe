from flask import Flask, jsonify, request
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_socketio import SocketIO, emit
import os
from .ai_service import analyze_resume_with_ai
import json
import fitz  # PyMuPDF
import docx  # python-docx
import io
import hashlib
import time
from backend.tasks import process_job_resumes
from datetime import datetime

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///resumes.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

CORS(app)  # This will enable CORS for all routes
socketio = SocketIO(app, cors_allowed_origins="*")

# --- Database Configuration ---
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'resumes.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# --- Database Models ---
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    jobs = db.relationship('Job', backref='user', lazy=True)

class Job(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    description = db.Column(db.Text, nullable=False)
    resumes = db.relationship('Resume', backref='job', lazy=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    candidate_name = db.Column(db.String(120), nullable=True)
    content = db.Column(db.Text, nullable=False)
    content_hash = db.Column(db.String(64), nullable=False, index=True)
    analysis = db.Column(db.Text, nullable=True)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)

    __table_args__ = (db.UniqueConstraint('job_id', 'filename', name='_job_filename_uc'),
                      db.UniqueConstraint('job_id', 'content_hash', name='_job_hash_uc'))

class Feedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    original_bucket = db.Column(db.String(50), nullable=False)
    suggested_bucket = db.Column(db.String(50), nullable=True)
    feedback_type = db.Column(db.String(20), nullable=False)  # 'override', 'correction', 'improvement'
    feedback_text = db.Column(db.Text, nullable=True)
    confidence_score = db.Column(db.Float, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    resume = db.relationship('Resume', backref='feedbacks')
    user = db.relationship('User', backref='feedbacks')

class BucketOverride(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    original_bucket = db.Column(db.String(50), nullable=False)
    new_bucket = db.Column(db.String(50), nullable=False)
    reason = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    resume = db.relationship('Resume', backref='bucket_overrides')
    user = db.relationship('User', backref='bucket_overrides')

# --- Interview Management Models ---
class Interview(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Interview Details
    title = db.Column(db.String(200), nullable=False)
    interview_type = db.Column(db.String(50), nullable=False)  # 'phone', 'video', 'onsite', 'technical'
    duration_minutes = db.Column(db.Integer, default=60)
    status = db.Column(db.String(20), default='scheduled')  # 'scheduled', 'completed', 'cancelled', 'rescheduled'
    
    # Scheduling
    scheduled_at = db.Column(db.DateTime, nullable=True)
    timezone = db.Column(db.String(50), default='UTC')
    location = db.Column(db.String(200), nullable=True)  # For onsite interviews
    video_link = db.Column(db.String(500), nullable=True)  # For video interviews
    
    # Interviewers
    primary_interviewer = db.Column(db.String(100), nullable=True)
    additional_interviewers = db.Column(db.Text, nullable=True)  # JSON array of interviewer names
    
    # Notes
    pre_interview_notes = db.Column(db.Text, nullable=True)
    post_interview_notes = db.Column(db.Text, nullable=True)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    resume = db.relationship('Resume', backref='interviews')
    job = db.relationship('Job', backref='interviews')
    user = db.relationship('User', backref='interviews')
    feedback = db.relationship('InterviewFeedback', backref='interview', uselist=False)

class InterviewFeedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    interview_id = db.Column(db.Integer, db.ForeignKey('interview.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Overall Assessment
    overall_rating = db.Column(db.Integer, nullable=False)  # 1-5 scale
    hire_recommendation = db.Column(db.String(20), nullable=False)  # 'strong_hire', 'hire', 'weak_hire', 'no_hire', 'strong_no_hire'
    
    # Detailed Ratings (JSON)
    technical_skills = db.Column(db.Integer, nullable=True)  # 1-5
    communication_skills = db.Column(db.Integer, nullable=True)  # 1-5
    problem_solving = db.Column(db.Integer, nullable=True)  # 1-5
    cultural_fit = db.Column(db.Integer, nullable=True)  # 1-5
    experience_relevance = db.Column(db.Integer, nullable=True)  # 1-5
    
    # Written Feedback
    strengths = db.Column(db.Text, nullable=True)
    areas_of_concern = db.Column(db.Text, nullable=True)
    additional_notes = db.Column(db.Text, nullable=True)
    
    # Interview Questions & Responses
    questions_asked = db.Column(db.Text, nullable=True)  # JSON array of questions
    candidate_responses = db.Column(db.Text, nullable=True)  # JSON array of responses
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    user = db.relationship('User', backref='interview_feedbacks')

class InterviewQuestion(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Question Details
    question_text = db.Column(db.Text, nullable=False)
    question_type = db.Column(db.String(50), nullable=False)  # 'technical', 'behavioral', 'situational', 'culture_fit'
    difficulty = db.Column(db.String(20), default='medium')  # 'easy', 'medium', 'hard'
    category = db.Column(db.String(100), nullable=True)  # e.g., 'algorithms', 'system_design', 'leadership'
    
    # Usage Tracking
    times_used = db.Column(db.Integer, default=0)
    avg_rating = db.Column(db.Float, default=0.0)  # Average rating from interviewers
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    job = db.relationship('Job', backref='interview_questions')
    user = db.relationship('User', backref='interview_questions')

# --- WebSocket Events ---
@socketio.on('connect')
def handle_connect():
    print('Client connected')
    emit('status', {'message': 'Connected to server'})

@socketio.on('disconnect')
def handle_disconnect():
    print('Client disconnected')

def emit_progress_update(job_id, message, progress_type='info', progress_data=None):
    """Emit progress updates to connected clients with optional progress data"""
    update_data = {
        'job_id': job_id,
        'message': message,
        'type': progress_type,
        'timestamp': time.time()
    }
    
    # Add progress data if provided
    if progress_data:
        update_data['progress'] = {
            'completed': progress_data.get('completed', 0),
            'total': progress_data.get('total', 0),
            'errors': progress_data.get('errors', 0),
            'percentage': round((progress_data.get('completed', 0) / max(progress_data.get('total', 1), 1)) * 100, 1)
        }
    
    socketio.emit('progress_update', update_data)

def check_job_completion(job_id):
    """Check if all resumes for a job are complete and emit completion event"""
    # This function is now handled by the progress tracking in tasks.py
    pass

def process_resume_with_progress(job_id, resume_file, job_description):
    """Process a single resume with progress updates"""
    try:
        emit_progress_update(job_id, f"Processing {resume_file.filename}...", 'processing')
        
        content = ""
        filename = resume_file.filename
        file_stream = resume_file.read()

        if filename.endswith('.pdf'):
            emit_progress_update(job_id, f"Reading PDF: {filename}", 'info')
            pdf_doc = fitz.open(stream=file_stream, filetype='pdf')
            for page in pdf_doc:
                content += page.get_text()
            pdf_doc.close()
        elif filename.endswith('.docx'):
            emit_progress_update(job_id, f"Reading DOCX: {filename}", 'info')
            doc = docx.Document(io.BytesIO(file_stream))
            for para in doc.paragraphs:
                content += para.text + '\n'
        else:
            content = file_stream.decode('utf-8')

        if not content.strip():
            emit_progress_update(job_id, f"Skipped {filename}: Empty or unreadable", 'warning')
            return None, 'File is empty or unreadable'

        # Check for duplicates
        content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
        existing_by_hash = Resume.query.filter_by(job_id=job_id, content_hash=content_hash).first()
        if existing_by_hash:
            emit_progress_update(job_id, f"Skipped {filename}: Duplicate content", 'warning')
            return None, 'Duplicate content'

        emit_progress_update(job_id, f"Analyzing {filename} with AI...", 'processing')
        analysis_text = analyze_resume_with_ai(job_description, content)
        analysis_json = json.loads(analysis_text)
        
        # Handle case where AI returns "Name Not Found"
        candidate_name = analysis_json.get('candidate_name', 'Not Provided')
        if candidate_name.strip().lower() == 'name not found':
            candidate_name = 'Not Provided'

        emit_progress_update(job_id, f"Completed analysis for {filename}", 'success')
        
        # Check if all resumes for this job are complete
        check_job_completion(job_id)
        
        return {
            'filename': filename,
            'candidate_name': candidate_name,
            'content': content,
            'content_hash': content_hash,
            'analysis': analysis_text
        }, None
        
    except Exception as e:
        emit_progress_update(job_id, f"Error processing {resume_file.filename}: {str(e)}", 'error')
        
        # Still check completion even on error
        check_job_completion(job_id)
        return None, str(e)

@app.route('/api/analyze', methods=['POST'])
def analyze_resumes():
    # --- Temp: Get or create a default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        default_user = User(username='default_user')
        db.session.add(default_user)
        db.session.commit()
    # --- End Temp ---

    if 'jobDescription' not in request.form:
        return jsonify({'error': 'No job description provided'}), 400

    job_description = request.form['jobDescription']
    resumes = request.files.getlist('resumes')

    # Check if a job with this description already exists FOR THIS USER.
    job = Job.query.filter_by(description=job_description, user_id=default_user.id).first()

    # If it doesn't exist, create a new one for this user.
    if not job:
        job = Job(description=job_description, user_id=default_user.id)
        db.session.add(job)
        db.session.flush()  # Use flush to get the job.id before committing.
        db.session.commit()  # Commit the job to the database

    emit_progress_update(job.id, f"Preparing {len(resumes)} resumes for background processing...", 'start')

    # Prepare resume data for background processing
    resumes_data = []
    for resume_file in resumes:
        if resume_file:
            try:
                content = ""
                filename = resume_file.filename
                file_stream = resume_file.read()

                if filename.endswith('.pdf'):
                    pdf_doc = fitz.open(stream=file_stream, filetype='pdf')
                    for page in pdf_doc:
                        content += page.get_text()
                    pdf_doc.close()
                elif filename.endswith('.docx'):
                    doc = docx.Document(io.BytesIO(file_stream))
                    for para in doc.paragraphs:
                        content += para.text + '\n'
                else:
                    content = file_stream.decode('utf-8')

                if content.strip():
                    resumes_data.append({
                        'filename': filename,
                        'content': content
                    })
                else:
                    emit_progress_update(job.id, f"Skipped {filename}: Empty or unreadable", 'warning')
                    
            except Exception as e:
                emit_progress_update(job.id, f"Error reading {resume_file.filename}: {str(e)}", 'error')

    if not resumes_data:
        emit_progress_update(job.id, "No valid resumes to process", 'warning')
        return jsonify({
            'message': 'No valid resumes to process',
            'job_id': job.id,
            'processed_files': [],
            'skipped_files': [{'filename': f.filename, 'reason': 'Invalid file'} for f in resumes if f]
        })

    # Initialize completion tracker for this job
    # Removed: job_completion_trackers[job.id] = {
    #     'total_resumes': len(resumes_data),
    #     'completed_resumes': 0
    # }

    # Queue background job using Celery
    process_job_resumes.delay(job.id, resumes_data, job_description)

    return jsonify({
        'message': f'Queued {len(resumes_data)} resumes for background processing',
        'job_id': job.id,
        'status': 'queued',
        'total_resumes': len(resumes_data)
    })

@app.route('/api/jobs', methods=['GET'])
def get_jobs():
    """Returns a list of all jobs for the current user."""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify([]) # No user, no jobs
    # --- End Temp ---

    jobs = Job.query.filter_by(user_id=default_user.id).order_by(Job.id.desc()).all()
    return jsonify([
        {
            'id': job.id, 
            'description': job.description,
            'resume_count': len(job.resumes)
        } for job in jobs
    ])

@app.route('/api/jobs/<int:job_id>', methods=['GET'])
def get_job_details(job_id):
    """Returns details and resumes for a specific job, checking user ownership."""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    job = Job.query.filter_by(id=job_id, user_id=default_user.id).first_or_404()
    
    resumes_data = [
        {
            'id': resume.id,
            'filename': resume.filename,
            'candidate_name': resume.candidate_name,
            'analysis': json.loads(resume.analysis) if resume.analysis else None
        } 
        for resume in job.resumes
    ]
    return jsonify({
        'id': job.id,
        'description': job.description,
        'resumes': resumes_data
    })

@app.route('/api/jobs/<int:job_id>', methods=['DELETE'])
def delete_job(job_id):
    """Delete a job and all its associated resumes"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    job = Job.query.filter_by(id=job_id, user_id=default_user.id).first()
    if not job:
        return jsonify({'error': 'Job not found'}), 404

    try:
        # Get resume count for confirmation
        resume_count = len(job.resumes)
        
        # Delete all associated resumes first (cascade)
        for resume in job.resumes:
            db.session.delete(resume)
        
        # Delete the job
        db.session.delete(job)
        db.session.commit()
        
        return jsonify({
            'message': f'Job "{job.description[:50]}..." deleted successfully',
            'deleted_resumes': resume_count,
            'job_id': job_id
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete job: {str(e)}'}), 500

@app.route('/api/data')
def get_data():
    return jsonify({'message': 'Hello from the Flask backend!'})

# --- Feedback Loop API Endpoints ---

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """Submit feedback for a resume analysis"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    required_fields = ['resume_id', 'original_bucket', 'feedback_type']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=data['resume_id']).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    # Check if resume belongs to user's job
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    try:
        feedback = Feedback(
            resume_id=data['resume_id'],
            user_id=default_user.id,
            original_bucket=data['original_bucket'],
            suggested_bucket=data.get('suggested_bucket'),
            feedback_type=data['feedback_type'],
            feedback_text=data.get('feedback_text'),
            confidence_score=data.get('confidence_score')
        )
        
        db.session.add(feedback)
        db.session.commit()
        
        return jsonify({
            'message': 'Feedback submitted successfully',
            'feedback_id': feedback.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to submit feedback: {str(e)}'}), 500

@app.route('/api/feedback/<int:resume_id>', methods=['GET'])
def get_resume_feedback(resume_id):
    """Get all feedback for a specific resume"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=resume_id).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    feedbacks = Feedback.query.filter_by(resume_id=resume_id).order_by(Feedback.created_at.desc()).all()
    
    return jsonify([{
        'id': f.id,
        'original_bucket': f.original_bucket,
        'suggested_bucket': f.suggested_bucket,
        'feedback_type': f.feedback_type,
        'feedback_text': f.feedback_text,
        'confidence_score': f.confidence_score,
        'created_at': f.created_at.isoformat()
    } for f in feedbacks])

@app.route('/api/override', methods=['POST'])
def submit_override():
    """Submit a bucket override for a resume"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    required_fields = ['resume_id', 'original_bucket', 'new_bucket', 'reason']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Verify resume exists
    resume = Resume.query.get(data['resume_id'])
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    try:
        # Create override record
        override = BucketOverride(
            resume_id=data['resume_id'],
            user_id=default_user.id,  # --- Temp: Use default user ---
            original_bucket=data['original_bucket'],
            new_bucket=data['new_bucket'],
            reason=data['reason'],
            timestamp=datetime.utcnow()
        )
        
        db.session.add(override)
        db.session.commit()
        
        return jsonify({'message': 'Override submitted successfully', 'override_id': override.id}), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to save override: {str(e)}'}), 500

@app.route('/api/override/<int:resume_id>', methods=['GET'])
def get_resume_overrides(resume_id):
    """Get all bucket overrides for a specific resume"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=resume_id).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    overrides = BucketOverride.query.filter_by(resume_id=resume_id).order_by(BucketOverride.created_at.desc()).all()
    
    return jsonify([{
        'id': o.id,
        'original_bucket': o.original_bucket,
        'new_bucket': o.new_bucket,
        'reason': o.reason,
        'created_at': o.created_at.isoformat()
    } for o in overrides])

@app.route('/api/feedback/stats', methods=['GET'])
def get_feedback_stats():
    """Get feedback statistics for the current user"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    try:
        # Get total feedback count
        total_feedback = Feedback.query.filter_by(user_id=default_user.id).count()
        
        # Get feedback by type
        feedback_by_type = db.session.query(
            Feedback.feedback_type,
            db.func.count(Feedback.id)
        ).filter_by(user_id=default_user.id).group_by(Feedback.feedback_type).all()
        
        # Get bucket override count
        total_overrides = BucketOverride.query.filter_by(user_id=default_user.id).count()
        
        # Get most common original buckets that get overridden
        common_overrides = db.session.query(
            BucketOverride.original_bucket,
            db.func.count(BucketOverride.id)
        ).filter_by(user_id=default_user.id).group_by(BucketOverride.original_bucket).order_by(
            db.func.count(BucketOverride.id).desc()
        ).limit(5).all()
        
        return jsonify({
            'total_feedback': total_feedback,
            'feedback_by_type': dict(feedback_by_type),
            'total_overrides': total_overrides,
            'common_overrides': dict(common_overrides)
        })
        
    except Exception as e:
        return jsonify({'error': f'Failed to get feedback stats: {str(e)}'}), 500

@app.route('/')
def home():
    return "Hello from the Backend!"

# --- Interview Management API Endpoints ---

@app.route('/api/interviews', methods=['GET'])
def get_interviews():
    """Get all interviews for the current user"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    try:
        interviews = Interview.query.filter_by(user_id=default_user.id).order_by(Interview.created_at.desc()).all()
        
        return jsonify([{
            'id': i.id,
            'resume_id': i.resume_id,
            'job_id': i.job_id,
            'title': i.title,
            'interview_type': i.interview_type,
            'duration_minutes': i.duration_minutes,
            'status': i.status,
            'scheduled_at': i.scheduled_at.isoformat() if i.scheduled_at else None,
            'timezone': i.timezone,
            'location': i.location,
            'video_link': i.video_link,
            'primary_interviewer': i.primary_interviewer,
            'additional_interviewers': json.loads(i.additional_interviewers) if i.additional_interviewers else [],
            'candidate_name': i.resume.candidate_name,
            'job_title': i.job.description[:100] + '...' if len(i.job.description) > 100 else i.job.description,
            'created_at': i.created_at.isoformat(),
            'updated_at': i.updated_at.isoformat()
        } for i in interviews])
        
    except Exception as e:
        return jsonify({'error': f'Failed to get interviews: {str(e)}'}), 500

@app.route('/api/interviews/<int:resume_id>', methods=['GET'])
def get_resume_interviews(resume_id):
    """Get all interviews for a specific resume"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=resume_id).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    try:
        interviews = Interview.query.filter_by(resume_id=resume_id).order_by(Interview.created_at.desc()).all()
        
        return jsonify([{
            'id': i.id,
            'title': i.title,
            'interview_type': i.interview_type,
            'duration_minutes': i.duration_minutes,
            'status': i.status,
            'scheduled_at': i.scheduled_at.isoformat() if i.scheduled_at else None,
            'timezone': i.timezone,
            'location': i.location,
            'video_link': i.video_link,
            'primary_interviewer': i.primary_interviewer,
            'additional_interviewers': json.loads(i.additional_interviewers) if i.additional_interviewers else [],
            'pre_interview_notes': i.pre_interview_notes,
            'post_interview_notes': i.post_interview_notes,
            'created_at': i.created_at.isoformat(),
            'updated_at': i.updated_at.isoformat(),
            'has_feedback': i.feedback is not None
        } for i in interviews])
        
    except Exception as e:
        return jsonify({'error': f'Failed to get interviews: {str(e)}'}), 500

@app.route('/api/interviews', methods=['POST'])
def create_interview():
    """Create a new interview"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    required_fields = ['resume_id', 'title', 'interview_type']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=data['resume_id']).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    try:
        # Parse scheduled_at if provided
        scheduled_at = None
        if data.get('scheduled_at'):
            scheduled_at = datetime.fromisoformat(data['scheduled_at'].replace('Z', '+00:00'))
        
        interview = Interview(
            resume_id=data['resume_id'],
            job_id=resume.job_id,
            user_id=default_user.id,
            title=data['title'],
            interview_type=data['interview_type'],
            duration_minutes=data.get('duration_minutes', 60),
            status=data.get('status', 'scheduled'),
            scheduled_at=scheduled_at,
            timezone=data.get('timezone', 'UTC'),
            location=data.get('location'),
            video_link=data.get('video_link'),
            primary_interviewer=data.get('primary_interviewer'),
            additional_interviewers=json.dumps(data.get('additional_interviewers', [])),
            pre_interview_notes=data.get('pre_interview_notes')
        )
        
        db.session.add(interview)
        db.session.commit()
        
        return jsonify({
            'message': 'Interview created successfully',
            'interview_id': interview.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to create interview: {str(e)}'}), 500

@app.route('/api/interviews/<int:interview_id>', methods=['PUT'])
def update_interview(interview_id):
    """Update an existing interview"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    interview = Interview.query.filter_by(id=interview_id, user_id=default_user.id).first()
    if not interview:
        return jsonify({'error': 'Interview not found'}), 404
    
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    try:
        # Update fields if provided
        if 'title' in data:
            interview.title = data['title']
        if 'interview_type' in data:
            interview.interview_type = data['interview_type']
        if 'duration_minutes' in data:
            interview.duration_minutes = data['duration_minutes']
        if 'status' in data:
            interview.status = data['status']
        if 'scheduled_at' in data:
            interview.scheduled_at = datetime.fromisoformat(data['scheduled_at'].replace('Z', '+00:00')) if data['scheduled_at'] else None
        if 'timezone' in data:
            interview.timezone = data['timezone']
        if 'location' in data:
            interview.location = data['location']
        if 'video_link' in data:
            interview.video_link = data['video_link']
        if 'primary_interviewer' in data:
            interview.primary_interviewer = data['primary_interviewer']
        if 'additional_interviewers' in data:
            interview.additional_interviewers = json.dumps(data['additional_interviewers'])
        if 'pre_interview_notes' in data:
            interview.pre_interview_notes = data['pre_interview_notes']
        if 'post_interview_notes' in data:
            interview.post_interview_notes = data['post_interview_notes']
        
        interview.updated_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({'message': 'Interview updated successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update interview: {str(e)}'}), 500

@app.route('/api/interviews/<int:interview_id>', methods=['DELETE'])
def delete_interview(interview_id):
    """Delete an interview"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    interview = Interview.query.filter_by(id=interview_id, user_id=default_user.id).first()
    if not interview:
        return jsonify({'error': 'Interview not found'}), 404
    
    try:
        db.session.delete(interview)
        db.session.commit()
        
        return jsonify({'message': 'Interview deleted successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete interview: {str(e)}'}), 500

@app.route('/api/interviews/<int:interview_id>/feedback', methods=['POST'])
def submit_interview_feedback(interview_id):
    """Submit feedback for an interview"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    interview = Interview.query.filter_by(id=interview_id, user_id=default_user.id).first()
    if not interview:
        return jsonify({'error': 'Interview not found'}), 404
    
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    required_fields = ['overall_rating', 'hire_recommendation']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    try:
        # Check if feedback already exists
        existing_feedback = InterviewFeedback.query.filter_by(interview_id=interview_id).first()
        if existing_feedback:
            return jsonify({'error': 'Feedback already exists for this interview'}), 400
        
        feedback = InterviewFeedback(
            interview_id=interview_id,
            user_id=default_user.id,
            overall_rating=data['overall_rating'],
            hire_recommendation=data['hire_recommendation'],
            technical_skills=data.get('technical_skills'),
            communication_skills=data.get('communication_skills'),
            problem_solving=data.get('problem_solving'),
            cultural_fit=data.get('cultural_fit'),
            experience_relevance=data.get('experience_relevance'),
            strengths=data.get('strengths'),
            areas_of_concern=data.get('areas_of_concern'),
            additional_notes=data.get('additional_notes'),
            questions_asked=json.dumps(data.get('questions_asked', [])),
            candidate_responses=json.dumps(data.get('candidate_responses', []))
        )
        
        db.session.add(feedback)
        db.session.commit()
        
        return jsonify({
            'message': 'Interview feedback submitted successfully',
            'feedback_id': feedback.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to submit feedback: {str(e)}'}), 500

@app.route('/api/interviews/<int:interview_id>/feedback', methods=['GET'])
def get_interview_feedback(interview_id):
    """Get feedback for an interview"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    interview = Interview.query.filter_by(id=interview_id, user_id=default_user.id).first()
    if not interview:
        return jsonify({'error': 'Interview not found'}), 404
    
    feedback = InterviewFeedback.query.filter_by(interview_id=interview_id).first()
    if not feedback:
        return jsonify({'error': 'No feedback found for this interview'}), 404
    
    return jsonify({
        'id': feedback.id,
        'overall_rating': feedback.overall_rating,
        'hire_recommendation': feedback.hire_recommendation,
        'technical_skills': feedback.technical_skills,
        'communication_skills': feedback.communication_skills,
        'problem_solving': feedback.problem_solving,
        'cultural_fit': feedback.cultural_fit,
        'experience_relevance': feedback.experience_relevance,
        'strengths': feedback.strengths,
        'areas_of_concern': feedback.areas_of_concern,
        'additional_notes': feedback.additional_notes,
        'questions_asked': json.loads(feedback.questions_asked) if feedback.questions_asked else [],
        'candidate_responses': json.loads(feedback.candidate_responses) if feedback.candidate_responses else [],
        'created_at': feedback.created_at.isoformat(),
        'updated_at': feedback.updated_at.isoformat()
    })

@app.route('/api/interviews/<int:interview_id>/feedback', methods=['PUT'])
def update_interview_feedback(interview_id):
    """Update feedback for an interview"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    interview = Interview.query.filter_by(id=interview_id, user_id=default_user.id).first()
    if not interview:
        return jsonify({'error': 'Interview not found'}), 404
    
    feedback = InterviewFeedback.query.filter_by(interview_id=interview_id).first()
    if not feedback:
        return jsonify({'error': 'No feedback found for this interview'}), 404
    
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    try:
        # Update fields if provided
        if 'overall_rating' in data:
            feedback.overall_rating = data['overall_rating']
        if 'hire_recommendation' in data:
            feedback.hire_recommendation = data['hire_recommendation']
        if 'technical_skills' in data:
            feedback.technical_skills = data['technical_skills']
        if 'communication_skills' in data:
            feedback.communication_skills = data['communication_skills']
        if 'problem_solving' in data:
            feedback.problem_solving = data['problem_solving']
        if 'cultural_fit' in data:
            feedback.cultural_fit = data['cultural_fit']
        if 'experience_relevance' in data:
            feedback.experience_relevance = data['experience_relevance']
        if 'strengths' in data:
            feedback.strengths = data['strengths']
        if 'areas_of_concern' in data:
            feedback.areas_of_concern = data['areas_of_concern']
        if 'additional_notes' in data:
            feedback.additional_notes = data['additional_notes']
        if 'questions_asked' in data:
            feedback.questions_asked = json.dumps(data['questions_asked'])
        if 'candidate_responses' in data:
            feedback.candidate_responses = json.dumps(data['candidate_responses'])
        
        feedback.updated_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({'message': 'Interview feedback updated successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to update feedback: {str(e)}'}), 500

@app.route('/api/interviews/questions/<int:job_id>', methods=['GET'])
def get_interview_questions(job_id):
    """Get interview questions for a job"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    # Verify job exists and belongs to user
    job = Job.query.filter_by(id=job_id, user_id=default_user.id).first()
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    
    try:
        questions = InterviewQuestion.query.filter_by(job_id=job_id).order_by(InterviewQuestion.created_at.desc()).all()
        
        return jsonify([{
            'id': q.id,
            'question_text': q.question_text,
            'question_type': q.question_type,
            'difficulty': q.difficulty,
            'category': q.category,
            'times_used': q.times_used,
            'avg_rating': q.avg_rating,
            'created_at': q.created_at.isoformat()
        } for q in questions])
        
    except Exception as e:
        return jsonify({'error': f'Failed to get questions: {str(e)}'}), 500

@app.route('/api/interviews/questions', methods=['POST'])
def create_interview_question():
    """Create a new interview question"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    required_fields = ['job_id', 'question_text', 'question_type']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Verify job exists and belongs to user
    job = Job.query.filter_by(id=data['job_id'], user_id=default_user.id).first()
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    
    try:
        question = InterviewQuestion(
            job_id=data['job_id'],
            user_id=default_user.id,
            question_text=data['question_text'],
            question_type=data['question_type'],
            difficulty=data.get('difficulty', 'medium'),
            category=data.get('category')
        )
        
        db.session.add(question)
        db.session.commit()
        
        return jsonify({
            'message': 'Interview question created successfully',
            'question_id': question.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to create question: {str(e)}'}), 500

@app.route('/api/interviews/questions/<int:question_id>', methods=['DELETE'])
def delete_interview_question(question_id):
    """Delete an interview question"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    question = InterviewQuestion.query.filter_by(id=question_id, user_id=default_user.id).first()
    if not question:
        return jsonify({'error': 'Question not found'}), 404
    
    try:
        db.session.delete(question)
        db.session.commit()
        
        return jsonify({'message': 'Interview question deleted successfully'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete question: {str(e)}'}), 500 